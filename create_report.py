import ged_lib as gl
import os
import docx
import cv2
import datetime
import operator
from docx import Document
from docx.shared import Cm
from docx.shared import Mm
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

def create_report():
#    print('\nCreating report')
    gl.get_params()

    if set_page_attributes (gl.page_size) == 0:
        return(0)

    read_image_types()
    gl.read_individuals()
    gl.read_families()
    gl.read_children()
    families_to_report.clear()
    
    if gl.initial_family == '0':
        read_families_to_report()
    else:
        family_id = gl.initial_family
        family_number = 1
        tree_walk(family_number, family_id)
        write_families_to_report()

    document = Document()
    sections = document.sections
    set_section_attributes(document, 1, 'P')
    
    if gl.title_page_required == 'Y':
        write_title_page(document, sections)
    
    if gl.contents_required == 'Y':
        write_contents(document)
    
    page_no = 1
    for i in range(0,len(families_to_report)):
        family_number = families_to_report[i].family_number
        family_id = families_to_report[i].family_id
        to_report = families_to_report[i].to_report

        add_family_to_index(family_id)
        if to_report == 'Y':
            write_family_page(document, family_number, family_id, page_no)
            page_no = page_no + 1
 
    if gl.index_required == 'Y': write_index(document)
    
    document.save('Report.docx')
    return(1)

def tree_walk(family_number, family_id):
    if family_id > 0:
        add_family_to_report(family_number, family_id, 'Y')   
        if gl.families[family_id].husband_id != 0:
            husband_id = gl.families[family_id].husband_id
            family_where_child = gl.get_family_where_child(husband_id)
            if family_where_child != 0:
                family_number = tree_walk(family_number + 1, family_where_child)

        if gl.families[family_id].wife_id != 0:
            wife_id = gl.families[family_id].wife_id
            family_where_child = gl.get_family_where_child(wife_id)
            if family_where_child != 0:
                family_number = tree_walk(family_number + 1, family_where_child)
    return(family_number)

def add_family_to_index(family_id):
    husband_id = gl.families[family_id].husband_id
    update_index(husband_id)
    wife_id = gl.families[family_id].wife_id
    update_index(wife_id)
    for i in range(0,len(gl.children)):
        child_family_id = gl.children[i].family_id
        if child_family_id != 0:
            if child_family_id > family_id: break
            if family_id == child_family_id:
                child_id = gl.children[i].child_id
                update_index(child_id)

def update_index(id):
    if id != 0:
        forename = gl.individuals[id].forename
        surname = gl.individuals[id].surname
        surname = surname.upper()
        name = forename
        if forename != '' and surname != '':
            name = name + ' '
        name = name + surname
        birth_date = gl.individuals[id].birth_date
        year = birth_date[-4:]
        key = surname + ', ' + forename + ' b. ' + year + ' ' + str(id)

        bookmark_family_number_where_spouse = ''
        family_number_where_spouse = 0
        family_id_where_spouse = gl.individuals[id].family_where_spouse
        if family_id_where_spouse != 0:
            family_number_where_spouse = get_family_number(family_id_where_spouse)
            if family_number_where_spouse > 0:
                bookmark_family_number_where_spouse = '[F' + str(family_number_where_spouse) + ']'

        bookmark_family_number_where_child = ''
        family_id_where_child = gl.individuals[id].family_where_child
        if family_id_where_child != 0:
            family_number_where_child = get_family_number(family_id_where_child)
            if family_number_where_child > 0:
                bookmark_family_number_where_child = '[F' + str(family_number_where_child) + ']'

        i = find_in_index(key)
        if i == -1:
            add_index(key, surname, forename, year, bookmark_family_number_where_spouse, \
                                        bookmark_family_number_where_child)
        else:
            index[i].family_number_where_spouse = bookmark_family_number_where_spouse
            index[i].family_number_where_child = bookmark_family_number_where_child

def find_in_index(key):
    r = len(index)
    for i in range(0, r):
        if index[i].key == key:
            return (i)
    return(-1)

def write_family_page(document, family_number, family_id, page_no):
    if gl.title_page_required == 'Y' or gl.contents_required == 'Y' or page_no > 1:
        document.add_page_break()
        section = new_section(document, 'P')
    else:
        section = document.sections[0]

    row = 0
    s_family_number = 'F' + str(family_number)
    bookmark = '[' + s_family_number + ']'
#    print ('\nFamily', bookmark, '[' + str(family_id) + ']')

    table = document.add_table(rows=1, cols=5)
    add_bookmark(table.rows[0].cells[0].paragraphs[0], bookmark)

    write_text(table, row, 2, 'Family ' + s_family_number)

    row = row + 1
    table.add_row()
    husband_id = gl.families[family_id].husband_id
    row = write_individual(table, 'Husband', row, husband_id)
        
    row = row + 1
    table.add_row()
    wife_id = gl.families[family_id].wife_id
    row = write_individual(table, 'Wife', row, wife_id)

    row = row + 1
    table.add_row()
    write_text(table, row, 0, 'Married')
    write_text(table, row, 3, gl.families[family_id].marriage_date)
    write_place(table, row, 4, gl.families[family_id].marriage_place)

    row = row + 1
    table.add_row()
    write_text(table, row, 0, 'Children')

    child_no = 0
    for i in range(0,len(gl.children)):
        children_family_id = gl.children[i].family_id
        if children_family_id != 0:
            if children_family_id > family_id: break
            if family_id == children_family_id:
                child_no = child_no + 1
                child_id = gl.children[i].child_id
                s_child_no = '#' + str(child_no)
                row = row + 1
                table.add_row()
                row = write_individual(table, s_child_no, row, child_id)
            
    set_col_widths(table)
    write_footer (section, ' ')

    if gl.document_images_required == 'Y':
        husband_id = families[family_id].husband_id
        if husband_id != 0:
            write_documents(document, husband_id)
        wife_id = families[family_id].wife_id
        if wife_id != 0:
            write_documents(document, wife_id)

def write_individual(table, individual_type, row, id):
    write_text(table, row, 0, individual_type)
    write_text(table, row, 2, 'Born')
    if id != 0:
        name = get_name_with_family_number(id, individual_type)
        if name[-1:] == ']':
            write_link(table, row, 1, name)
        else:
            write_text(table, row, 1, name)
        
        birth_date = gl.individuals[id].birth_date
        birth_place = gl.individuals[id].birth_place
        write_text(table, row, 3, birth_date)
        write_place(table, row, 4, birth_place)
        
        baptism_date = gl.individuals[id].baptism_date
        baptism_place = gl.individuals[id].baptism_place
        if baptism_date != '' or baptism_place != '':
            row = row + 1
            table.add_row()
            write_text(table, row, 2, 'Baptised')
            write_text(table, row, 3, baptism_date)
            write_place(table, row, 4, baptism_place)
    
        death_date = gl.individuals[id].death_date
        death_place = gl.individuals[id].death_place
        if death_date != '' or death_place != '':
            row = row + 1
            table.add_row()
            write_text(table, row, 2, 'Died')
            write_text(table, row, 3, death_date)
            write_place(table, row, 4, death_place)
    
        burial_date = gl.individuals[id].burial_date
        burial_place = gl.individuals[id].burial_place
        if burial_date != '' or burial_place != '':
            row = row + 1
            table.add_row()
            write_text(table, row, 2, 'Buried')
            write_text(table, row, 3, burial_date)
            write_place(table, row, 4, burial_place)
        
    return(row)

def write_text(table, row, col, text):
    table.cell(row, col).text = text
    table.cell(row, col).paragraphs[0].runs[0].font.name = font_name
    table.cell(row, col).paragraphs[0].runs[0].font.size = Pt(normal_font_size)
    if row == 0:
        table.cell(row, col).paragraphs[0].runs[0].font.bold = True
        table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def write_link(table, row, col, text):
    paragraph = table.cell(row, col).paragraphs[0]
    link_to = text[text.index('['):]
    add_link (paragraph, link_to, text)

def write_place(table, row, col, text):
    text_to_remove = ', ' + gl.country_to_remove1
    text = text.replace(text_to_remove,'')
    text_to_remove = ', ' + gl.country_to_remove2
    text = text.replace(text_to_remove,'')
    write_text(table, row, col, text)

def set_col_widths(table):
    widths = (Cm(1.80), Cm(5.40), Cm(2.40), Cm(2.48), Cm(4.48))
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

def add_bookmark(paragraph, bookmark_name):
    run = paragraph.add_run()
    tag = run._r  

    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)
    
def add_link(paragraph, link_to, text):
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    run = paragraph.add_run()
    run._r.append (hyperlink)
    run.font.name = font_name
    run.font.size = Pt(normal_font_size)
    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True

def write_documents(document, id):
    forename = individuals[id].forename
    name = individuals[id].surname
    name = name + ', ' + forename
    birth_date = individuals[id].birth_date
    year = birth_date[-4:]
    name = name + ' b. ' + year
    path = gl.website_path + "/images"
    for image_type in it.image_types:
        file_to_find = name + ' - ' + image_type + '.jpg'
        for file in os.listdir(path):
            if file.lower() == file_to_find.lower():
                full_file_name = path + '\\' + file
#                print (full_file_name) #debug
                img = cv2.imread(full_file_name)
                height, width = img.shape[:2]
                if height < width:
                    section = new_section(document, 'L')
                    image_ratio = width / height
                    if image_ratio < max_image_ratio:
                        height = max_image_width
                        document.add_picture(full_file_name, height = Mm(height))
                    else:                
                        width = max_image_height
                        document.add_picture(full_file_name, width = Mm(width))
                else:
                    section = new_section(document, 'P')
                    image_ratio = height / width
                    if image_ratio < max_image_ratio:
                        width = max_image_width
                        document.add_picture(full_file_name, width = Mm(width))
                    else:                
                        height = max_image_height
                        document.add_picture(full_file_name, height = Mm(height))
                
                last_paragraph = document.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                text = file.replace('.jpg','')
                text = text.replace('BIR', 'Birth Index')
                text = text.replace('MIR', 'Marriage Index')
                text = text.replace('DIR', 'Death Index')
                write_footer (section, text)
            
def write_footer (section, text):
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].font.bold = True
    paragraph.runs[0].font.name = font_name
    paragraph.runs[0].font.size = Pt(normal_font_size)

def write_title_page(document, sections):
    section = sections[0]
    table = document.add_table(rows=title_page_rows, cols=1)
    write_title_text(table, 5, 0, gl.title1)
    write_title_text(table, 6, 0, gl.title2)
    write_title_text(table, 7, 0, gl.title3)
    write_title_text(table, 8, 0, gl.title4)
    write_text(table, title_page_rows - 2, 0, 'Prepared on')
    today = datetime.date.today()
    x = datetime.datetime(today.year, today.month, today.day)
    write_text(table, title_page_rows - 1, 0, x.strftime('%d %b %Y'))
    write_footer (section, ' ')

def write_title_text(table, row, col, text):
    table.cell(row, col).text = text
    table.cell(row, col).paragraphs[0].runs[0].font.name = font_name
    table.cell(row, col).paragraphs[0].runs[0].font.size = Pt(title_font_size)
    table.cell(row, col).paragraphs[0].runs[0].font.bold = True
    table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def write_contents(document):
    if gl.title_page_required == 'Y':
        document.add_page_break()
        section = new_section(document, 'P')
    else:
        section = document.sections[0]
        
    table = document.add_table(rows=len(families_to_report) + 1, cols=1)
    write_text (table, 0, 0, 'Contents')
    for i in range(0,len(families_to_report)):
        husband_name = ''
        wife_name = ''
        family_number = families_to_report[i].family_number
        family_id = families_to_report[i].family_id
        if family_id == 0: break

        husband_id = families[family_id].husband_id
        if husband_id != 0:
            husband_name = get_name(husband_id)
        
        wife_id = families[family_id].wife_id
        if wife_id != 0:
            wife_name = get_name(wife_id)
        
        s_family_number = str(family_number)
        contents_record = 'F' + s_family_number + '   ' + husband_name + ', ' + wife_name
        paragraph = table.cell(i + 1, 0).paragraphs[0]
        link_to = '[F' + s_family_number + ']' 
        add_link (paragraph, link_to, contents_record)
    write_footer (section, ' ')

def write_index(document):
    section = new_section(document, 'P')
    idx.index.sort(key=operator.attrgetter('key'))
    r = len(idx.index)
    for i in range(0, r):
        index_within_page = i % (index_rows * 2) # remainder
        if index_within_page == 0: # if first record on a page
            if i > 0: # if not the first index page
                write_footer (section, ' ')
#                document.add_page_break()
            table = document.add_table(rows=1, cols=1)
            write_text (table, 0, 0, '    Index')
            table = document.add_table(rows=index_rows, cols=2)
            last_surname = ''
        
        row_within_page = index_within_page % index_rows # remainder 
        col_within_page = index_within_page // index_rows # integer divide 
        
        index_record = idx.index[i].surname
        if index_record == last_surname and index_record != '' and row_within_page > 0:
            index_record = '-----'
        else:
            last_surname = index_record
        index_record = index_record + ', ' + idx.index[i].forename + ' '
        year = idx.index[i].birth_year
        if year != '':
            index_record = index_record + 'b. ' + year + ' '
        
        table.cell(row_within_page, col_within_page).text = index_record
        table.cell(row_within_page, col_within_page).paragraphs[0].runs[0].font.name = font_name
        table.cell(row_within_page, col_within_page).paragraphs[0].runs[0].font.size = Pt(normal_font_size)

        family_number_where_spouse = idx.index[i].family_number_where_spouse
        family_number_where_child = idx.index[i].family_number_where_child
        if family_number_where_spouse != 0:
            add_link(table.cell(row_within_page, col_within_page).paragraphs[0], \
                     family_number_where_spouse, family_number_where_spouse)
            if family_number_where_child != 0: 
                run = table.cell(row_within_page, col_within_page).paragraphs[0].add_run()
                run.text = ', '
                run.font.name = font_name
                run.font.size = Pt(normal_font_size)

        if family_number_where_child != 0:
            add_link(table.cell(row_within_page, col_within_page).paragraphs[0], \
                     family_number_where_child, family_number_where_child)
    write_footer (section, ' ')

def set_page_attributes(size):
    global page_size, font_name, page_height, page_width, left_margin, right_margin, top_margin, bottom_margin, \
           header_len, footer_len, title_font_size, title_page_rows, normal_font_size, index_rows, \
           max_image_height, max_image_width, max_image_ratio
    
    page_size = size
    font_name = "Calibri"
    if page_size == "A4":
        page_height = 297
        page_width = 210
        left_margin = 15
        right_margin = 15
        top_margin = 15
        bottom_margin = 15
        header_len = 10
        footer_len = 10
        title_font_size = 26
        title_page_rows = 26
        normal_font_size = 11
        index_rows = 28
        max_image_height = page_height - top_margin - bottom_margin - header_len - footer_len - 10
        max_image_width = page_width - left_margin - right_margin - 10
        max_image_ratio = max_image_height / max_image_width
        return(1)
    if page_size == "A5":
        page_height = 210
        page_width = 148
        left_margin = 10
        right_margin = 10
        top_margin = 10
        bottom_margin = 10
        header_len = 6
        footer_len = 6
        title_font_size = 18
        title_page_rows = 18
        normal_font_size = 8
        index_rows = 20
        max_image_height = page_height - top_margin - bottom_margin - header_len - footer_len - 10
        max_image_width = page_width - left_margin - right_margin - 10
        max_image_ratio = max_image_height / max_image_width
        return(1)
    print ("Unsupported Page Size")
    return(0)
 
def set_section_attributes(document, n, orientation):
    section = document.sections[n-1]
    if orientation == 'L':
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(page_height)
        section.page_height = Mm(page_width)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_height = Mm(page_height)
        section.page_width = Mm(page_width)

    section.left_margin = Mm(left_margin)
    section.right_margin = Mm(right_margin)
    section.top_margin = Mm(top_margin)
    section.bottom_margin = Mm(bottom_margin)
    section.header_distance = Mm(header_len)
    section.footer_distance = Mm(footer_len)

def new_section(document, orientation):
    section = document.add_section()
    set_section_attributes(document, len(document.sections), orientation)
    return(section)

class families_to_report_class(object):
    def __init__(self, family_number=None, family_id=None, to_report=None):
        self.family_number = family_number
        self.family_id = family_id
        self.to_report = to_report

families_to_report = []

def read_families_to_report():
    file = open('FamiliesToReport.txt','r')
    families_to_report.clear()
    while True:
        s = file.readline()
        s = s.strip()
        if s == '':
            break
        x = s.split("~")
        add_family_to_report(int(x[0]), int(x[1]), x[6])

def add_family_to_report(family_number, family_id, to_report):
    families_to_report.append(families_to_report_class(family_number, family_id, to_report))
    
def write_families_to_report():
    file = open('FamiliesToReport.txt','w')
    for i in range(0,len(families_to_report)):
        family_id = int(families_to_report[i].family_id)
        husbands_name = ""
        husbands_birth_date = ""
        if gl.families[family_id].husband_id != "":
            husband_id = int(gl.families[family_id].husband_id)
            if husband_id > 0:
                husbands_name = gl.get_person_name(husband_id)
                husbands_birth_date = gl.get_birth_year(husband_id)

        wifes_name = ""
        wifes_birth_date = ""
        if gl.families[family_id].wife_id != "":
            wife_id = int(gl.families[family_id].wife_id)
            if wife_id > 0:
                wifes_name = gl.get_person_name(wife_id)
                wifes_birth_date = gl.get_birth_year(wife_id)

        line = str(families_to_report[i].family_number) + "~"
        s = str(family_id)
        line = line + s + "~"
        line = line + husbands_name + "~"
        line = line + husbands_birth_date + "~"
        line = line + wifes_name + "~"
        line = line + wifes_birth_date + "~"
        line = line + families_to_report[i].to_report
        file.write(line + '\n')
    
    file.close()

def get_family_number(family_id):
    if (family_id != 0):
        for i in range(0,len(families_to_report)):
            if family_id == families_to_report[i].family_id:
                return(families_to_report[i].family_number)
    return(0)
    
def get_name_with_family_number (person_id, person_type):
    name = gl.get_name(person_id)
    if person_type == 'Husband' or person_type == 'Wife':
        family_id = gl.individuals[person_id].family_where_child
    else:
        family_id = gl.individuals[person_id].family_where_spouse

    family_number = get_family_number(family_id)

    if family_number > 0:
        name = name + ' [F' + str(family_number) + ']'
    return(name)

class index_class(object):
    def __init__(self, key=None, surname=None, forename=None, birth_year=None, family_number_where_spouse=None, \
                 family_number_where_child=None):
        self.key = key
        self.surname = surname
        self.forename = forename
        self.birth_year = birth_year
        self.family_number_where_spouse = family_number_where_spouse
        self.family_number_where_child = family_number_where_child
        
index = []

def add_index(key, surname, forename, birth_year, family_number_where_spouse, family_number_where_child):
    index.append(index_class(key, surname, forename, birth_year, family_number_where_spouse,
                             family_number_where_child))
    
image_types = []

def read_image_types():
    global image_types

    imagetypefile = open("ImageTypes.txt","r")
    while True:
        s = imagetypefile.readline()
        s = s.strip()
        if s == "":
            break
        image_types.append(s)
        
page_size = ""
font_name = ""
page_height = 0
page_width = 0
left_margin = 0
right_margin = 0
top_margin = 0
bottom_margin = 0
header_len = 0
footer_len = 0
title_font_size = 0
title_page_rows = 0
normal_font_size = 0
index_rows = 0
max_image_height = 0
max_image_width = 0
max_image_ratio = 0

#create_report()
